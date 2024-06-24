import os
import sys
import nltk
import pandas as pd
import requests
from bs4 import BeautifulSoup
from pymongo import MongoClient, UpdateOne, errors
from concurrent.futures import ThreadPoolExecutor, as_completed
from cachetools import cached, TTLCache
import logging
from urllib3 import Retry
from requests.adapters import HTTPAdapter
from langdetect import detect, DetectorFactory
from openai import OpenAI

# Initialize logging
logging.basicConfig(level=logging.INFO)

# MongoDB Connection
def get_mongo_connection():
    try:
        mongo_uri = 'mongodb+srv://doadmin:6d30Bi4ec59u7ag1@egypt-horizon-scanner-1948d167.mongo.ondigitalocean.com/egypt-horizon-scanner?tls=true&authSource=admin&replicaSet=egypt-horizon-scanner'
        client = MongoClient(mongo_uri, tls=True, authSource='admin')
        db = client["egypt-horizon-scanner"]
        client.admin.command('ping')  # Check connection
        return db["emergence_issue_of_the_month_data"]
    except errors.ServerSelectionTimeoutError as err:
        logging.error("Server selection timeout error: %s", err)
        sys.exit(1)
    except errors.ConnectionFailure as err:
        logging.error("Connection failure: %s", err)
        sys.exit(1)
    except Exception as err:
        logging.error("An unexpected error occurred: %s", err)
        sys.exit(1)

# %pip install pandoc python-docx openpyxl PyMuPDF transformers mammoth pywin32 Spire.Doc
# %pip install doc2docx pytesseract
# %pip install textract aspose-words
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import os
import json
from docx import Document
import fitz  # PyMuPDF
from openpyxl import load_workbook
import aspose.words as aw
import spacy
import pytesseract
from PIL import Image
from openai import OpenAI

# Load a pre-trained NER model
nlp = spacy.load("en_core_web_sm")
# List of questions to extract specific fields
questions = {
    "id": "What is the project ID?",
    "projectNo": "What is the project number?",
    "projectName": "What is the name of the project?",
    "projectDetail": "Can you provide a detailed description of the project?",
    "photoURL": "What is the photo URL associated with the project?", 
    "executingAgency": "Who is the executing agency for the project?",
    "status": "What is the status of the project? Is it completed, in progress, or not started?",
    "theme": "What are the main themes or focus areas of the project?",
    "estimatedCost": "What is the estimated cost of the project?",
    "budget": "What is the budget allocation for the project?",
    "totalDonatedAmount": "What is the total amount donated to the project?",
    "startDate": "What is the start date of the project? When did the project begin? Write the date in the format DD/MM/YYYY.",
    "endDate": "What is the end date of the project? When is the project expected to be completed? Write the date in the format DD/MM/YYYY.",
    # "Latitude": "What are the latitude coordinates of the project location?",
    # "Longitude": "What are the longitude coordinates of the project location?",
    # "Locality_Name_EN": "What is the locality name in English?",
    # "Locality_Name_AR": "What is the locality name in Arabic? Translation: ما هو اسم المنطقة بالعربية؟",
    # "Locality_PCODE": "What is the locality postal code?",
    # "City_Name_EN": "What is the city name in English?",
    # "City_Name_AR": "What is the city name in Arabic? Translation: ما هو اسم المدينة بالعربية؟",
    # "City_PCODE": "What is the city postal code?",
    # "District_Name_EN": "What is the district name in English?",
    # "District_Name_AR": "What is the district name in Arabic? Translation: ما هو اسم الحي بالعربية؟",
    # "District_PCODE": "What is the district postal code?",
    "Municipal_Division_Type": "What is the municipal division type (kism, markaz, new city, or police-administered)?",
    "Municipal_Division_Name_EN": "What is the municipal division name in English?",
    "Municipal_Division_Name_AR": "What is the municipal division name in Arabic? Translation: ما هو اسم القسم البلدي بالعربية؟",
    "Municipal_Division_PCODE": "What is the municipal division postal code?",
    "Governorate_Name_EN": "What is the governorate name in English?",
    "Governorate_Name_AR": "What is the governorate name in Arabic? Translation: ما هو اسم المحافظة بالعربية؟",
    "Governorate_PCODE": "What is the governorate postal code?",
    # "State_Name_EN": "What is the state name in English?",
    # "State_Name_AR": "What is the state name in Arabic? Translation: ما هو اسم الولاية بالعربية؟",
    # "State_PCODE": "What is the state postal code?",
    # "Province_Name_EN": "What is the province name in English?",
    # "Province_Name_AR": "What is the province name in Arabic? Translation: ما هو اسم المحافظة بالعربية؟",
    # "Province_PCODE": "What is the province postal code?",
    # "Region_Name_EN": "What is the region name in English?",
    # "Region_Name_AR": "What is the region name in Arabic? Translation: ما هو اسم المنطقة بالعربية؟",
    # "Region_PCODE": "What is the region postal code?",
    "Country_EN": "What is the country name in English?",
    "Country_AR": "What is the country name in Arabic?",
    "Country_PCODE": "What is the country postal code? Translation: ما هو الرمز البريدي للبلد؟",
    "donor": "Who are the donors for the project?",
    "contribution": "What contributions have been made to the project?",
    "dataReliability": "What is the data reliability rating for the project information?"
}
# List of questions to extract specific fields
answers = {
    "id": "[project ID]",
    "projectNo": "[Project number]",
    "projectName": "[Project Name]",
    "projectDetail": "[Project Detail]",
    "photoURL": "[Photo URL]", 
    "executingAgency": "[Executing Agency]",
    "status": "[Status: Not Started, In Progress, Completed]",
    "theme": "[Theme]",
    "estimatedCost": "[Estimated Cost]",
    "budget": "[Budget]",
    "totalDonatedAmount": "[Total Donated Amount]",
    "startDate": "[Start Date]",
    "endDate": "[End Date]",
    # "Latitude": "[Latitude]",
    # "Longitude": "[Longitude]",
    # "Locality_Name_EN": "[Locality Name EN]",
    # "Locality_Name_AR": "[Locality Name AR]",
    # "Locality_PCODE": "[Locality PCODE]",
    # "City_Name_EN": "[City Name EN]",
    # "City_Name_AR": "[City Name AR]",
    # "City_PCODE": "[City PCODE]",
    # "District_Name_EN": "[District Name EN]",
    # "District_Name_AR": "[District Name AR]",
    # "District_PCODE": "[District PCODE]",
    "Municipal_Division_Type": "[Municipal Division Type]",
    "Municipal_Division_Name_EN": "[Municipal Division Name EN]",
    "Municipal_Division_Name_AR": "[Municipal Division Name AR]",
    "Municipal_Division_PCODE": "[Municipal Division PCODE]",
    "Governorate_Name_EN": "[Governorate Name EN]",
    "Governorate_Name_AR": "[Governorate Name AR]",
    "Governorate_PCODE": "[Governorate PCODE]",
    # "State_Name_EN": "[State Name EN]",
    # "State_Name_AR": "[State Name AR]",
    # "State_PCODE": "[State PCODE]",
    # "Province_Name_EN": "[Province Name EN]",
    # "Province_Name_AR": "[Province Name AR]",
    # "Province_PCODE": "[Province PCODE]",
    # "Region_Name_EN": "[Region Name EN]",
    # "Region_Name_AR": "[Region Name AR]",
    # "Region_PCODE": "[Region PCODE]",
    "Country_EN": "[Country Name EN]",
    "Country_AR": "[Country Name AR]",
    "Country_PCODE": "[Country PCODE]",
    "donor": "[Donor]",
    "contribution": "[Contribution]",
    "dataReliability": "[Data Reliability]"
    
}
	1.	id: 1827
	2.	projectNo: PIMS 1237 CC MSP
	3.	projectName: Introduction of Vehicle Electric Bus Technology and Hybrid-Electric Bus Technology in Egypt – Phase 1a
	4.	projectDetail:
		•	The project aimed to introduce electric and hybrid-electric buses in Egypt to reduce pollution, improve public health, and protect national monuments from degradation due to heavy-duty buses.
		•	Phase 1a involved testing two electric buses, training local staff on maintenance and operations, and planning for subsequent phases.
	5.	photoURL: Not provided
	6.	executingAgency: [Egyptian Environmental Affairs Agency (EEAA), Social Fund for Development (SFD)]
	7.	status: Completed (Phase 1a)
	8.	theme: [Climate Change, Sustainable Transportation, Public Health, Cultural Heritage Protection]
	9.	estimatedCost: USD 1.714 million
	10.	budget:
		•	GEF: USD 0.7486 million
		•	EEAA: USD 0.3154 million
		•	SFD: USD 0.1 million
		•	In-kind contribution: USD 0.55 million
	11.	totalDonatedAmount: USD 0.7486 million
	12.	startDate: March 2000
	13.	endDate: June 2006 (Phase 1a)
	14.	Latitude: Not provided
	15.	Longitude: Not provided
	16.	Locality_Name_EN: Giza, Luxor
	17.	Locality_Name_AR: الجيزة، الأقصر
	18.	Locality_PCODE: Not provided
	19.	City_Name_EN: Cairo
	20.	City_Name_AR: القاهرة
	21.	City_PCODE: Not provided
	22.	District_Name_EN: Not provided
	23.	District_Name_AR: Not provided
	24.	District_PCODE: Not provided
	25.	Municipal_Division_Type: Not provided
	26.	Municipal_Division_Name_EN: Not provided
	27.	Municipal_Division_Name_AR: Not provided
	28.	Municipal_Division_PCODE: Not provided
	29.	Governorate_Name_EN: Giza Governorate, Luxor Governorate
	30.	Governorate_Name_AR: محافظة الجيزة، محافظة الأقصر
	31.	Governorate_PCODE: Not provided
	32.	State_Name_EN: Not applicable
	33.	State_Name_AR: Not applicable
	34.	State_PCODE: Not applicable
	35.	Province_Name_EN: Not applicable
	36.	Province_Name_AR: Not applicable
	37.	Province_PCODE: Not applicable
	38.	Region_Name_EN: Not applicable
	39.	Region_Name_AR: Not applicable
	40.	Region_PCODE: Not applicable
	41.	Country_EN: Egypt
	42.	Country_AR: مصر
	43.	Country_PCODE: Not provided
	44.	donor: [Global Environment Facility (GEF)]
	45.	contribution: USD 0.7486 million (GEF), USD 0.3154 million (EEAA), USD 0.1 million (SFD)
	46.	dataReliability: Information derived from final evaluation reports and project documents.
client = OpenAI(api_key="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")
# Function to get response from GPT-3.5-turbo
def gpt_get(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0,
            #response_format={ "type": "json_object" }
        )
    return response.choices[0].message.content.strip(), response.usage.prompt_tokens, response.usage.completion_tokens
# Function to extract information
def extract_info(text, questions):
    extracted_info = {}
    max_tokens = 15000  # Adjusted to avoid exceeding token limit 16384
    text_chunks = [text[i:i+max_tokens] for i in range(0, len(text), max_tokens)]
    
    for key, question in questions.items():
        extracted_info[key] = []
        for chunk in text_chunks:
            prompt = f"Question: {question}\nContext: {chunk}\nPlease provide the information in the following format, the answer should be short:\n\n" \
                     f"{answers[key]}\n\nAnswer:"
            answer, _, _ = gpt_get(prompt)
            extracted_info[key].append(answer)

    return extracted_info
# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text
# Function to extract text from XLSX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text
# Function to convert DOC to DOCX and then extract text
def convert_doc_to_docx(doc_path):
    docx_path = doc_path + "x"
    doc = aw.Document(doc_path)
    doc.save(docx_path)
    return docx_path
# Function to extract text from DOC
def extract_text_from_doc(doc_path):
    docx_path = convert_doc_to_docx(doc_path)
    text = extract_text_from_docx(docx_path)
    
    os.remove(docx_path)
    return text
# Function to extract text from XLSX
def extract_data_from_xlsx(xlsx_path):
    wb = load_workbook(xlsx_path)
    ws = wb.active
    data = "\n".join(["\t".join(map(str, row)) for row in ws.iter_rows(values_only=True)])
    return data
# Main function to extract text based on file type
def extract_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.doc'):
        return extract_text_from_doc(file_path)
    elif file_path.endswith('.xlsx'):
        return extract_data_from_xlsx(file_path)
    else:
        raise ValueError("Unsupported file format")
# Function to extract named entities from text using a specific model
def extract_entities(text):
    model_name = "dbmdz/bert-large-cased-finetuned-conll03-english"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForTokenClassification.from_pretrained(model_name)
    nlp = pipeline("ner", model=model, tokenizer=tokenizer, grouped_entities=True)
    entities = nlp(text)
    return entities
# Function to extract text from an image
def extract_text_from_image(image_path):
    return pytesseract.image_to_string(Image.open(image_path))
# Function to extract entities from text
def extract_entities(text):
    doc = nlp(text)
    return [(ent.text, ent.label_) for ent in doc.ents]
# Function to extract information from files and save to JSON
def extract_info_from_files(file_paths, json_path):
    texts = [extract_text_from_file(file_path) for file_path in file_paths]
    combined_text = "\n".join(texts)
    
    # entities = extract_entities(combined_text)
    # Extract the information
    extracted_information = extract_info(combined_text, questions)
    # project_data = map_entities_to_schema(extracted_information)

    with open(json_path, 'w') as json_file:
        json.dump(extracted_information, json_file, indent=4)
file_paths = [
"/Users/izzymohamed/Downloads/Documents/PIMS 1237 CC MSP - ELECTRIC BUSES/Electric_buses_eval-Executive_Summary.doc",
"/Users/izzymohamed/Downloads/Documents/PIMS 1237 CC MSP - ELECTRIC BUSES/Electric_buses_Final_Evaluation_report.doc",
"/Users/izzymohamed/Downloads/Documents/PIMS 1237 CC MSP - ELECTRIC BUSES/TERMS_OF_REFERENCE.doc",
"/Users/izzymohamed/Downloads/Documents/PIMS 1237 CC MSP - ELECTRIC BUSES/report.xlsx"
]

json_path = "output.json"



if __name__ == "__main__":
    try:
        db_collection = get_mongo_connection()
        
        # Fetching all data from the collection
        data = list(db_collection.find())
        
        logging.info(f"START: Total records: {len(data)}")
        
        # Validate URLs
        valid_data = [item for item in data if isinstance(item, dict) and item.get('link')]
        
        if not valid_data:
            logging.error("No valid URLs found in the input data.")
            sys.exit(1)
        
        # Attempt to process URLs
        update_emerging_issues_data(db_collection)
    except Exception as e:
        logging.error("An error occurred: %s", e)
        sys.exit(1)