# Import necessary libraries
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import os
import json
from docx import Document
import fitz  # PyMuPDF
from openpyxl import load_workbook
import aspose.words as aw
import pytesseract
from PIL import Image
from openai import OpenAI
import spacy
import logging

# Initialize logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Load a pre-trained NER model
nlp = spacy.load("en_core_web_sm")

# Define the questions and answers schema
questions = {
    "id": "What is the project ID?",
    "projectNo": "What is the project number?",
    "projectName": "What is the name of the project?",
    "projectDetail": "Can you provide a detailed description of the project?",
    "photoURL": "What is the photo URL associated with the project?", 
    "executingAgency": "Who is the executing agency for the project?",
    "status": "What is the status of the project? Is it completed, in progress, or not started?",
    "theme": "What are the main themes or focus areas of the project?",
    "estimatedCost": "What is the estimated cost of the project?",
    "budget": "What is the budget allocation for the project?",
    "totalDonatedAmount": "What is the total amount donated to the project?",
    "startDate": "What is the start date of the project? When did the project begin? Write the date in the format DD/MM/YYYY.",
    "endDate": "What is the end date of the project? When is the project expected to be completed? Write the date in the format DD/MM/YYYY.",
    "Municipal_Division_Type": "What is the municipal division type (kism, markaz, new city, or police-administered)?",
    "Municipal_Division_Name_EN": "What is the municipal division name in English?",
    "Municipal_Division_Name_AR": "What is the municipal division name in Arabic? Translation: ما هو اسم القسم البلدي بالعربية؟",
    "Municipal_Division_PCODE": "What is the municipal division postal code?",
    "Governorate_Name_EN": "What is the governorate name in English?",
    "Governorate_Name_AR": "What is the governorate name in Arabic? Translation: ما هو اسم المحافظة بالعربية؟",
    "Governorate_PCODE": "What is the governorate postal code?",
    "Country_EN": "What is the country name in English?",
    "Country_AR": "What is the country name in Arabic?",
    "Country_PCODE": "What is the country postal code? Translation: ما هو الرمز البريدي للبلد؟",
    "donor": "Who are the donors for the project?",
    "contribution": "What contributions have been made to the project?",
    "dataReliability": "What is the data reliability rating for the project information?"
}

# Amended answers schema
answers = {
    "id": "[project ID]",
    "projectNo": "[Project number]",
    "projectName": "[Project Name]",
    "projectDetail": "[Project Detail]",
    "photoURL": "[Photo URL]",
    "executingAgency": "[Executing Agency]",
    "status": "[Status: Not Started, In Progress, Completed]",
    "theme": "[Theme]",
    "estimatedCost": "[Estimated Cost]",
    "budget": "[Budget]",
    "totalDonatedAmount": "[Total Donated Amount]",
    "startDate": "[Start Date: DD/MM/YYYY]",
    "endDate": "[End Date: DD/MM/YYYY]",
    "Municipal_Division_Type": "[Municipal Division Type]",
    "Municipal_Division_Name_EN": "[Municipal Division Name EN]",
    "Municipal_Division_Name_AR": "[Municipal Division Name AR]",
    "Municipal_Division_PCODE": "[Municipal Division PCODE]",
    "Governorate_Name_EN": "[Governorate Name EN]",
    "Governorate_Name_AR": "[Governorate Name AR]",
    "Governorate_PCODE": "[Governorate PCODE]",
    "Country_EN": "[Country Name EN]",
    "Country_AR": "[Country Name AR]",
    "Country_PCODE": "[Country PCODE]",
    "donor": "[Donor]",
    "contribution": "[Contribution]",
    "dataReliability": "[Data Reliability]"
}

# Initialize OpenAI client
client = OpenAI(api_key="your_openai_api_key_here")

# Function to get response from GPT-3.5-turbo
def gpt_get(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message.content.strip(), response.usage.prompt_tokens, response.usage.completion_tokens

# Function to extract information
def extract_info(text, questions):
    extracted_info = {}
    max_tokens = 15000  # Adjusted to avoid exceeding token limit 16384
    text_chunks = [text[i:i+max_tokens] for i in range(0, len(text), max_tokens)]
    
    for key, question in questions.items():
        extracted_info[key] = []
        for chunk in text_chunks:
            prompt = f"Question: {question}\nContext: {chunk}\nPlease provide the information in the following format, the answer should be short. Give one final answer:\n\n{answers[key]}\n\nAnswer:"
            answer, _, _ = gpt_get(prompt)
            extracted_info[key].append(answer)

    return extracted_info

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

# Function to convert DOC to DOCX and then extract text
def convert_doc_to_docx(doc_path):
    docx_path = doc_path + "x"
    doc = aw.Document(doc_path)
    doc.save(docx_path)
    return docx_path

# Function to extract text from DOC
def extract_text_from_doc(doc_path):
    docx_path = convert_doc_to_docx(doc_path)
    text = extract_text_from_docx(docx_path)
    os.remove(docx_path)
    return text

# Function to extract text from XLSX
def extract_data_from_xlsx(xlsx_path):
    wb = load_workbook(xlsx_path)
    ws = wb.active
    data = "\n".join(["\t".join(map(str, row)) for row in ws.iter_rows(values_only=True)])
    return data

# Main function to extract text based on file type
def extract_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.doc'):
        return extract_text_from_doc(file_path)
    elif file_path.endswith('.xlsx'):
        return extract_data_from_xlsx(file_path)
    else:
        raise ValueError("Unsupported file format")

# Function to extract text from an image
def extract_text_from_image(image_path):
    return pytesseract.image_to_string(Image.open(image_path))

# Function to extract named entities from text using a specific model
def extract_entities(text):
    model_name = "dbmdz/bert-large-cased-finetuned-conll03-english"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForTokenClassification.from_pretrained(model_name)
    nlp = pipeline("ner", model=model, tokenizer=tokenizer, grouped_entities=True)
    entities = nlp(text)
    return entities

# Function to extract information from files and save to JSON
def extract_info_from_files(file_paths, json_path):
    texts = [extract_text_from_file(file_path) for file_path in file_paths]
    combined_text = "\n".join(texts)
    
    # Extract the information
    extracted_information = extract_info(combined_text, questions)

    with open(json_path, 'w') as json_file:
        json.dump(extracted_information, json_file, indent=4)

if __name__ == "__main__":
    file_paths = [
        "/path/to/Electric_buses_eval-Executive_Summary.doc",
        "/path/to/Electric_buses_Final_Evaluation_report.doc",
        "/path/to/TERMS_OF_REFERENCE.doc",
        "/path/to/report.xlsx"
    ]
    
    json_path = "output.json"
    extract_info_from_files(file_paths, json_path)
    logging.info("Extraction complete. Results saved to %s", json_path)